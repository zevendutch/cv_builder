from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

# profile pic
document.add_picture('ZevenDutch.png', width=Inches(3.0))

# personal details
# name = 'ZevenDutch'
# phone_number = '000000'
# email = 'zevendutch@gmail.com'

new_name = input('What is your name? ')
new_phone_number = input('What is your phone number? ')
new_email = input('What is your email? ')

speak('Hello ' + new_name + ' how are you today?')
speak('Your email is interesting, I am confirming it is ' + new_email)

# document.add_paragraph(
#     name + ' | ' + phone_number + ' | ' + email)

document.add_paragraph(
    new_name + ' | ' + new_phone_number + ' | ' + new_email)

# about me
document.add_heading('About me')
about_me = input('Tell about yourself ')
document.add_paragraph(about_me)

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From date ')
to_date = input('To date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company + ' ')

p.add_run(experience_details)


# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From date ')
        to_date = input('To date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ' ')

        p.add_run(experience_details)

    else:
        break


# skills
document.add_heading('Skills')
p = document.add_paragraph()

skill = input('Enter skill ')
years = input('Years of Experience ')

p.style = 'List Bullet'

p.add_run(skill + ' ').bold = True
p.add_run(' - ' + years + ' years')

while True:
    has_more_skills = input(
        'Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        p = document.add_paragraph()
        skill = input('Enter skill ')
        years = input('Years of Experience ')
        p.add_run(skill + ' ').bold = True
        p.add_run(' - ' + years + ' years')
        p.style = 'List Bullet'

    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Amigoscode tutorial on YouTube"

document.save('cv.docx')
